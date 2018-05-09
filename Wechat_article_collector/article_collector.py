from datetime import *
import wechatsogou
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from add_hyperlinks import add_hyperlink


# 文章爬取
def get_articles(headline=True, original=True, timedel=1, add_account=None):

    with open('gzh.txt', 'r') as f:
        accounts = [account.strip() for account in f.readlines()]
    # add_account必须是一个list或None
    if add_account is not None:
        if isinstance(list, add_account):
            accounts.extend(add_account)
            with open('gzh.txt', 'w') as f:
                for account in accounts:
                    f.write(account)
        else:
            print('add_account should be a list')

    ws_api = wechatsogou.WechatSogouAPI(captcha_break_time=3)
    articles = []
    for account in accounts:
        articles.extend(reformat(ws_api.get_gzh_article_by_history(account)))

    # 时间过滤，只选取规定天数以内的
    timestamp = int((datetime.now()-timedelta(days=timedel)).timestamp())
    articles = [article for article in articles if article['datetime'] > timestamp]

    # 头条文章过滤，是否选取头条文章，默认是
    if headline:
        articles = [article for article in articles if article['main'] == 1]

    # 原创文章过滤，是否选取原创文章，默认是
    if original:
        articles = [article for article in articles if article['copyright_stat'] == 100]

    return articles


# 为保存每篇文章的字典添加一个公众号来源
def reformat(data):
    atcs = data.get('article')
    if atcs is not None:
        wechat_name = data.get('gzh')['wechat_name']
        for article in atcs:
            article['wechat_name'] = wechat_name
        return atcs


# 文章整合为文本
def to_msdocx(data):
    document = Document()
    header = '公众号最新文章({})'.format(datetime.now().strftime('%a, %b %d %H:%M'))
    document.add_heading(header, 0)
    for article in data:
        document.add_paragraph(article['title'], style='ListNumber')
        document.add_paragraph('摘要： ' + article['abstract'])
        p = document.add_paragraph('链接： ')
        add_hyperlink(p, '小鼠标点击这里哦', article['content_url'])
        document.add_paragraph('来自： ' + article['wechat_name']+'\n')
    p = document.add_paragraph('今天也要元气满满哦~')
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_picture('比心.JPG', width=Inches(1))
    document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.save('/Users/apple/Desktop/{}.docx'.format(header))


if __name__ == "__main__":
    articles = get_articles(timedel=1)
    to_msdocx(articles)
